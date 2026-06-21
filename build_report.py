"""
Generateur du rapport academique (.docx) - projet IDS Intelligent / dataset ORNL Driver ID.
Retrace tout le parcours : pieges methodologiques -> modeles -> evaluation -> signature
d'injection -> robustesse -> auto-critique. Adapte du generateur de l'ancien projet.
Sortie : deliverables/Rapport_IDS_Intelligent.docx       Lancer : python build_report.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(ROOT, "docs", "assets")
OUT = os.path.join(ROOT, "deliverables", "Rapport_IDS_Intelligent.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()
stl = doc.styles["Normal"]; stl.font.name = "Calibri"; stl.font.size = Pt(11)
for h, sz in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
    s = doc.styles[h]; s.font.name = "Calibri"; s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
NAVY = RGBColor(0x1F, 0x3A, 0x5F); GREY = "EEF2F7"


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), color); tcPr.append(shd)


def shade_para(p, color):
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), color); pPr.append(shd)


def H(text, level=1): doc.add_heading(text, level)


def P(text, bold=False, italic=False, size=None, align=None, color=None):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p


def bullets(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")


def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")


def callout(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + " - "); r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10)
    r2 = p.add_run(text); r2.font.size = Pt(10)
    shade_para(p, "EAF1F8")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""; r = hdr[i].paragraphs[0].add_run(h); r.bold = True
        r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(hdr[i], "1F3A5F")
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; rr = cells[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(10)
            if ridx % 2: shade(cells[i], GREY)
    doc.add_paragraph(); return t


def figure(name, caption, width=6.0):
    path = os.path.join(ASSETS, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        P(caption, italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        P(f"[figure manquante : {name}]", italic=True)


def code(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(9)
    shade_para(p, "F4F6F8")


def pagebreak(): doc.add_page_break()


def add_toc():
    p = doc.add_paragraph(); run = p.add_run()
    a = OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"), "begin")
    b = OxmlElement("w:instrText"); b.set(qn("xml:space"), "preserve"); b.text = 'TOC \\o "1-3" \\h \\z \\u'
    c = OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"), "separate")
    d = OxmlElement("w:t"); d.text = "Clic droit ici -> 'Mettre a jour les champs' pour generer la table des matieres."
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (a, b, c, d, e): run._r.append(el)


# ============================================================================ PAGE DE GARDE
for _ in range(3): doc.add_paragraph()
P("Detecteur d'intrusion intelligent pour bus CAN de poids lourd", bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
P("Detection par machine learning d'une cyberattaque sur le reseau embarque", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
P("(CAN / J1939) d'un camion Kenworth T270 - dataset ORNL Driver Identification", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
P("Rapport detaille - parcours d'apprentissage", italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(2): doc.add_paragraph()
P("Projet integrateur 2026", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
P("Encadrant : Rida Khatoun", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(5): doc.add_paragraph()
P("Jeu de donnees : ORNL Driver Identification (Kenworth T270, 50 conducteurs)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
P("155 902 fenetres de 1 s - attaque rare (1,46 %)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
pagebreak()

H("Table des matieres", 1); add_toc(); pagebreak()

# ============================================================================ RESUME
H("Resume", 1)
P("Ce rapport retrace la conception complete d'un systeme de detection d'intrusion (IDS) "
  "pour le reseau embarque d'un camion, a partir du dataset ORNL Driver Identification : une "
  "flotte de 50 conducteurs sur un Kenworth T270 dont l'enregistreur electronique (ELD) a ete "
  "compromis pour injecter une cyberattaque sur l'afficheur. Au-dela de la consigne - construire "
  "et comparer des modeles de machine learning distinguant trafic normal et malveillant - le "
  "projet est mene comme un parcours rigoureux : chaque notion est expliquee, appliquee, mesuree, "
  "puis critiquee.")
P("Resultats cles : (1) la difficulte centrale n'est pas le modele mais les CONFONDEURS - "
  "l'attaque survient toujours au meme lieu/moment, ce qui piege un modele naif (un split "
  "aleatoire donne PR-AUC 0,985, trompeur, contre 0,632 par conducteur) ; (2) le champion "
  "Gradient Boosting sur les signaux CAN atteint PR-AUC 0,756 (0,798 apres optimisation), bien "
  "au-dessus du deep (0,54-0,57) et de l'anomalie (~0,02) ; (3) nous ISOLONS une signature de "
  "l'injection elle-meme - le silence du bus CAN0 ~4 s apres l'attaque - distincte de la reaction "
  "du conducteur ; (4) nous montrons honnetement les limites : detecteur evadable, mono-attaque, "
  "biometrie inutile. Un demonstrateur interactif accompagne le rapport.")
P("Le fil conducteur, en cinq temps :", bold=True)
numbered([
    "Comprendre les pieges : confondeurs de lieu/temps et fuite par conducteur - c'est le coeur "
    "methodologique, et la vraie difficulte du sujet.",
    "Comparer les familles de modeles (arbres, SVM, reseaux, anomalie) et designer un champion.",
    "Evaluer honnetement : courbe ROC vs precision-rappel, seuil operationnel, latence par episode.",
    "Isoler la signature de l'injection elle-meme (le silence du bus CAN0) - la decouverte du projet.",
    "Eprouver les limites (attaquant adaptatif, detecteur mono-attaque) et les assumer."])
callout("Comment lire ce rapport", "Les encarts bleus (Intuition, Pour comprendre, A retenir) "
        "condensent les idees essentielles. Le chapitre 3 pose la theorie ; les chapitres 5-6 "
        "exposent le coeur methodologique (confondeurs, fuite par conducteur).")
pagebreak()

# ============================================================================ 1. CONTEXTE
H("1. Introduction et contexte", 1)
H("1.1 Le terrain : un camion et son reseau interne", 2)
P("Un vehicule lourd moderne est un reseau de calculateurs (ECU) qui dialoguent sur un bus CAN "
  "(Controller Area Network). Au-dessus du CAN, le protocole J1939 normalise le sens de chaque "
  "message pour les poids lourds.")
callout("Pour comprendre le CAN", "Le bus CAN est un fil partage que tous les calculateurs "
        "ecoutent (diffusion). Un message ne porte pas d'adresse de destinataire mais un "
        "identifiant decrivant le type d'information. Concu pour la fiabilite temps reel, il n'a "
        "aucune securite : ni authentification, ni chiffrement. Quiconque accede au bus peut "
        "emettre n'importe quel message - c'est la faille exploitee ici.")
P("Chaque signal J1939 est designe par un SPN (Suspect Parameter Number). Le dataset les fournit "
  "deja decodes et agreges a la seconde (moyenne, ecart-type, min, max).")

H("1.2 L'attaque", 2)
P("Un ELD (Electronic Logging Device) est un boitier reglementaire branche sur le bus pour "
  "enregistrer les heures de conduite. Lors d'un test controle a l'ORNL, le firmware de l'ELD a "
  "ete remplace a distance par Wi-Fi depuis un vehicule roulant a proximite ; l'ELD compromis "
  "peut alors ecrire des messages CAN arbitraires. L'attaque enregistree cible l'afficheur "
  "(instrument cluster) : elle met a zero le tachymetre / compteur de vitesse.")
P("Particularite essentielle du protocole experimental : 50 conducteurs ont subi l'attaque, "
  "repartis en trois groupes d'AWARENESS (G1 : non prevenu ; G2 : prevenu ; G3 : prevenu et "
  "instruit de se garer). L'attaque survient toujours au MEME endroit du trajet et au MEME moment "
  "(~60 % du parcours). Ce dispositif est au coeur de toute la methodologie (chapitres 5-6).")

H("1.3 Le jeu de donnees", 2)
table(["Element", "Valeur"],
      [["Source", "ORNL Driver Identification Dataset (Kenworth T270)"],
       ["Unite", "fenetre de 1 s (mean / sd / min / max par signal)"],
       ["Fenetres", "155 902"],
       ["Conducteurs", "50 (3 groupes d'awareness)"],
       ["Modalites", "CAN J1939 (337) / biometrie HR-EDA (7) / GPS-inertie VBOX (312)"],
       ["Cible", "cyberattack_active - attaque RARE (1,46 %)"]])
pagebreak()

# ============================================================================ 2. PROBLEMATIQUE
H("2. Problematique et objectifs", 1)
H("2.1 La question centrale et sa vraie difficulte", 2)
P("La question litterale - distinguer automatiquement une fenetre normale d'une fenetre sous "
  "attaque - est facile a enoncer mais piegee. Car l'attaque etant toujours au meme lieu/moment, "
  "un modele peut 'detecter' l'endroit ou l'instant plutot que l'attaque elle-meme. Toute la "
  "difficulte est de ne pas se faire pieger par ces CONFONDEURS.")
H("2.2 Ce qui rend le probleme difficile", 2)
table(["Difficulte", "Consequence methodologique"],
      [["Attaque au meme LIEU", "exclure les features GPS/inertie (sinon on detecte l'endroit)"],
       ["Attaque au meme MOMENT", "surveiller les signaux a derive lente (confondeur temps)"],
       ["Un seul trajet par conducteur", "split PAR CONDUCTEUR (sinon fuite massive)"],
       ["Attaque rare (1,46 %)", "metrique PR-AUC, jamais l'accuracy"],
       ["Cible = injection + reaction", "interpreter ce qu'on detecte vraiment"]])
H("2.3 Objectifs (consigne)", 2)
numbered(["Explorer et caracteriser normal vs suspect.",
          "Pretraiter : nettoyage, normalisation, selection de features.",
          "Comparer plusieurs algorithmes (arbres, SVM, reseaux de neurones...).",
          "Decouper en entrainement / test (par conducteur, anti-fuite).",
          "Entrainer sur comportements normaux et malveillants.",
          "Evaluer (precision, rappel, courbe ROC) et optimiser les hyperparametres."])
pagebreak()

# ============================================================================ 3. THEORIE ML
H("3. Fondements theoriques du machine learning", 1)
P("Ce chapitre pose le socle d'apprentissage mobilise dans le projet. Une version autonome et "
  "plus detaillee est disponible dans docs/00_theorie/fondements.md.")
H("3.1 Qu'est-ce que le machine learning ?", 2)
P("La programmation classique applique des regles ecrites a la main : donnees + regles -> "
  "resultats. Le machine learning inverse la fleche : donnees + resultats -> regles. On montre "
  "des exemples (entrees et bonnes reponses) et le programme apprend une regle qui generalise.")
callout("Le risque central", "Le sur-apprentissage (overfitting) : un modele qui apprend par "
        "coeur le bruit du train et echoue sur des donnees nouvelles. On le controle en evaluant "
        "TOUJOURS sur des donnees non vues - ici, des conducteurs non vus a l'entrainement.")
H("3.2 Les types d'apprentissage", 2)
table(["Type", "Idee", "Labels ?", "Notre usage"],
      [["Supervise", "apprendre entree -> etiquette", "oui", "chemin A (champion)"],
       ["Non supervise", "trouver structure / anomalies", "non", "chemin B, clustering"],
       ["Semi-supervise", "peu d'etiquetes + beaucoup sans", "partiel", "self-training (V2)"],
       ["Deep / sequentiel", "apprendre la representation", "oui", "MLP, GRU"]])
H("3.3 Le workflow scikit-learn", 2)
table(["Etape", "Role", "Pourquoi"],
      [["Chargement", "structurer en nombres", "le modele ingere des nombres"],
       ["Split train/test", "separer apprendre et juger", "mesurer la generalisation"],
       ["Preparation", "imputer, normaliser", "memes echelles (sauf arbres)"],
       ["Modele", "choisir l'algorithme", "selon donnee et tache"],
       ["Tuning", "regler les hyperparametres", "extraire le meilleur du modele"],
       ["Evaluation", "predire et mesurer", "sur des donnees jamais vues"]])
callout("Normalisation", "Les modeles a distance/gradient (SVM, regression, reseaux) exigent des "
        "features a la meme echelle ; les arbres (Random Forest, Gradient Boosting) n'en ont pas "
        "besoin (ils comparent des seuils). Le scaler s'ajuste sur le train seul (anti-fuite).")
H("3.4 La matrice de confusion et les metriques", 2)
table(["", "Predit Normal", "Predit Attaque"],
      [["Reel Normal", "TN", "FP (fausse alerte)"],
       ["Reel Attaque", "FN (intrusion ratee)", "TP"]])
bullets(["Precision = TP/(TP+FP) : parmi les alertes, la part de vraies.",
         "Rappel = TP/(TP+FN) : parmi les attaques, la part detectee.",
         "ROC-AUC : separation tous seuils confondus (optimiste si positifs rares).",
         "PR-AUC : aire precision-rappel ; la metrique HONNETE quand l'attaque est rare."])
callout("ROC vs PR (capital ici)", "Avec 98,5 % de negatifs, la courbe ROC parait excellente "
        "(0,977) car les nombreux negatifs ecrasent les faux positifs. La PR-AUC (0,735) se "
        "concentre sur les positifs : c'est la bonne lecture en contexte rare.")
H("3.5 Validation et fuite de donnees", 2)
P("On apprend sur le train, on juge sur le test. Si des donnees quasi identiques sont des deux "
  "cotes, le test est contamine : c'est la fuite de donnees. Ici, chaque conducteur ne fait qu'un "
  "trajet ; un split aleatoire mettrait des fenetres du meme conducteur des deux cotes - le modele "
  "reconnaitrait la PERSONNE, pas l'attaque. D'ou le split PAR CONDUCTEUR (GroupKFold).")
pagebreak()

# ============================================================================ 4. METHODO
H("4. Methodologie generale", 1)
P("Chaque brique suit la meme boucle : theorie -> application au CAN -> experience mesuree -> "
  "verdict assume. Le projet s'organise en phases (cadrage, EDA, pretraitement, modelisation, "
  "evaluation) puis en VAGUES d'amelioration apres auto-critique :")
table(["Vague", "Objet", "Exemples"],
      [["Vague 1", "credibilite / conformite au sujet", "ROC, tuning, latence, multi-seed, tests"],
       ["Vague 2", "profondeur methodologique", "signature d'injection, evasion, taxonomie"]])
P("Trois chemins de modelisation sont compares (supervise, anomalie, deep), tous en validation "
  "PAR CONDUCTEUR, metrique PR-AUC.")
pagebreak()

# ============================================================================ 5. EDA
H("5. Exploration des donnees (EDA)", 1)
P("L'exploration n'est pas optionnelle : c'est elle qui revele les confondeurs et dicte toute la "
  "strategie.")
H("5.1 Le confondeur de LIEU (le piege GPS)", 2)
P("L'attaque est 7 a 44 fois plus concentree geographiquement que le trajet : elle survient "
  "toujours apres un point precis du parcours. Un modele utilisant le GPS detecte donc l'ENDROIT, "
  "pas l'attaque (Cohen's d GPS = 2,63). On EXCLUT les 312 features GPS/inertie (VBOX).")
figure("eda_confounder_gps.png", "Figure 5.1 - Concentration geographique de l'attaque (confondeur de lieu).", 5.5)
H("5.2 Le confondeur de TEMPS", 2)
P("L'attaque survient a ~60 % du trajet ; les signaux a derive lente (temperatures) correlent "
  "avec l'heure de roulage. On les surveille (test CAN_STABLE au chapitre 6).")
figure("eda_cohens_d.png", "Figure 5.2 - Pouvoir discriminant des signaux (Cohen's d).", 5.5)
H("5.3 La reponse biometrique (faible)", 2)
P("La frequence cardiaque monte pendant l'attaque (+1,7 a +3,1 bpm), modulee par le groupe "
  "d'awareness (max pour le groupe 2). Effet REEL mais modeste (d=0,14) - une piste, pas un "
  "detecteur (confirme au chapitre 11).")
figure("eda_biometric_hr.png", "Figure 5.3 - Reponse cardiaque pendant l'attaque, par groupe.", 5.0)
pagebreak()

# ============================================================================ 6. PRETRAITEMENT / SPLIT
H("6. Pretraitement et le coeur methodologique : le split par conducteur", 1)
P("Le dataset etant deja etiquete et decode, l'enjeu n'est pas le labeling mais d'eviter les "
  "fuites. C'est ici que se joue la credibilite de tout le projet.")
H("6.1 La fuite par conducteur, chiffree", 2)
P("On compare la meme tache (Gradient Boosting, features CAN) selon le decoupage :")
table(["Protocole de split", "PR-AUC", "Lecture"],
      [["Aleatoire (fenetres melangees)", "0,985", "TROMPEUR : fuite par conducteur"],
       ["Par conducteur (honnete)", "0,632", "la vraie difficulte"],
       ["Avec GPS, par conducteur", "0,835", "spurieux : geofencing du lieu"]])
callout("Pourquoi c'est decisif", "Un split aleatoire affiche 0,985 : le modele reconnait le "
        "conducteur (vu des deux cotes), pas l'attaque. En splittant PAR CONDUCTEUR on tombe a "
        "0,632 - le point de depart honnete. Et avec le GPS on remonte a 0,835, mais c'est le "
        "lieu, pas l'attaque. Ces deux pieges dictent nos garde-fous : split conducteur + CAN seul.")
figure("p2_confounders_demo.png", "Figure 6.1 - Les pieges chiffres (fuite conducteur, confondeur GPS).", 6.0)
H("6.2 Le test CAN_STABLE (confondeur temps)", 2)
P("En excluant les 21 signaux les plus correles au temps de roulage, la PR-AUC passe de 0,632 a "
  "0,630 : inchangee. Le signal CAN est donc GENUINE, il ne vient pas du confondeur temps.")
pagebreak()

# ============================================================================ 7. FEATURES
H("7. Features et modalites", 1)
P("On classe les colonnes en trois modalites et on mesure leur pouvoir discriminant (ablation, "
  "Gradient Boosting, split conducteur) :")
table(["Modalite", "Nb", "PR-AUC", "Verdict"],
      [["CAN (signaux J1939)", "337", "0,756", "modalite de travail"],
       ["CAN + biometrie", "344", "0,749", "la biometrie n'aide pas"],
       ["Biometrie seule", "7", "0,014", "= hasard"],
       ["GPS / inertie", "312", "0,835", "CONFONDEUR - exclu"]])
callout("A retenir", "Le GPS 'bat' le CAN (0,835 > 0,756) mais c'est un piege (le lieu). La "
        "biometrie seule est au niveau du hasard. On travaille sur le CAN seul (337 signaux).")
pagebreak()

# ============================================================================ 8. MODELISATION
H("8. Modelisation multi-chemins", 1)
H("8.1 Chemin A - apprentissage supervise", 2)
P("On compare plusieurs familles. Intuition de chacune :")
bullets(["Regression logistique : frontiere lineaire + probabilite (sigmoide). Rapide, limitee au lineaire.",
         "SVM : maximise la marge entre classes ; avec un noyau RBF, devient non lineaire mais couteuse.",
         "Random Forest : moyenne de nombreux arbres - reduit la variance, robuste.",
         "Gradient Boosting : arbres en sequence, chacun corrigeant le precedent. Souvent le meilleur en tabulaire."])
P("Resultats (PR-AUC, par conducteur ; hasard ~ 0,015) :")
table(["Modele", "PR-AUC"],
      [["Regression logistique", "0,407"],
       ["SVM lineaire", "0,390"],
       ["Random Forest", "0,709"],
       ["Gradient Boosting (CHAMPION)", "0,756 +/- 0,088"]])
callout("Pourquoi les arbres boostes gagnent", "L'attaque se distingue par des motifs non "
        "lineaires dans les signaux CAN. Une frontiere lineaire ne peut pas les capter ; un "
        "ensemble d'arbres, qui combine de nombreux seuils, oui.")
figure("p4a_supervised.png", "Figure 8.1 - Comparaison supervisee et ablation des modalites.", 6.0)
H("8.2 Chemin B - detection d'anomalie", 2)
P("On apprend uniquement le NORMAL (novelty detection) - interet : detecter une attaque inconnue. "
  "Isolation Forest, One-Class SVM, gaussienne, PCA. Resultat : les 4 detecteurs sont AU NIVEAU "
  "DU HASARD (~0,02). La variabilite inter-conducteur noie l'attaque ; rare n'est pas aberrant. "
  "Echec assume qui justifie l'approche supervisee.")
figure("p4b_anomaly.png", "Figure 8.2 - Detection d'anomalie : echec assume (~ hasard).", 5.5)
H("8.3 Chemin C - deep learning", 2)
P("MLP tabulaire et GRU temporel (fenetres de 16 s) en PyTorch (GPU). Resultats multi-graines : "
  "MLP 0,543 +/- 0,016 et GRU 0,571 +/- 0,024 - l'ecart est NON significatif et aucun ne bat les "
  "arbres (0,756).")
callout("Lecon importante", "Le deep ne bat pas le Gradient Boosting. Sur des donnees tabulaires "
        "deja agregees avec peu de positifs, les arbres boostes dominent. Le deep learning n'est "
        "pas un label de qualite : on choisit l'outil selon la donnee.")
figure("p4c_deep.png", "Figure 8.3 - Chemin C (deep) vs arbres.", 5.0)
pagebreak()

# ============================================================================ 9. EVALUATION
H("9. Evaluation du champion et optimisation", 1)
H("9.1 Courbes ROC et PR (la demande du sujet)", 2)
P("Predictions hors-fold (GroupKFold conducteur). AUC-ROC = 0,977 alors que PR-AUC = 0,735. Le "
  "contraste est pedagogique : la classe negative ecrase les faux positifs. On fournit la ROC "
  "(demandee) mais on s'appuie sur la PR-AUC (attaque rare).")
figure("v1_roc_vs_pr.png", "Figure 9.1 - ROC (0,977) vs precision-rappel (0,735).", 5.5)
H("9.2 Optimisation des hyperparametres (la demande du sujet)", 2)
P("RandomizedSearchCV (40 candidats), meme validation par conducteur. Le champion par defaut "
  "passe de PR-AUC 0,757 a 0,798 optimise (+0,040). Gain reel mais modeste (inferieur a "
  "l'ecart-type inter-fold) : sur un Gradient Boosting deja solide, le tuning aide a la marge. "
  "Meilleurs parametres : arbres peu profonds (max_leaf_nodes=15), learning_rate=0,1, "
  "l2_regularization=10,0.")
H("9.3 Metriques IDS : latence et detection par episode", 2)
P("Un IDS doit detecter l'EPISODE, pas la fenetre. Sur les 51 episodes d'attaque (duree mediane "
  "42 s), au seuil F1-max : 86 % d'episodes detectes, latence mediane 4 s, 0,24 % de fausses "
  "alertes. Bien meilleur que ne le suggere le score par fenetre.")
figure("v1_latency_episode.png", "Figure 9.2 - Detection par episode et latence.", 5.5)
H("9.4 Generalisation par conducteur (LODO bimodale)", 2)
P("Leave-one-driver-out (50 conducteurs) : mediane 0,920, mais 7 conducteurs < 0,50 dont 6 du "
  "Groupe 1. La generalisation est BIMODALE : excellente sauf pour les conducteurs non avertis.")
figure("p5b_group_analysis.png", "Figure 9.3 - Gradient d'awareness : detectabilite G1 < G2 < G3.", 6.0)
callout("Le gradient d'awareness", "La detectabilite suit l'awareness : G1 (non averti) 0,74 / "
        "G2 0,92 / G3 0,96. L'IDS detecte en partie la REACTION du conducteur (G3 se gare -> le "
        "regime chute), pas seulement l'injection. Cette limite est elucidee au chapitre 10.")
pagebreak()

# ============================================================================ 10. SIGNATURE INJECTION
H("10. La decouverte : isoler la signature d'injection", 1)
P("La faille de validite la plus grave : la cible melange l'INJECTION CAN et la REACTION du "
  "conducteur. On cherche une signature de l'injection elle-meme, independante de la reaction.")
H("10.1 Le bus CAN0 se tait pendant l'attaque", 2)
P("Le regime moteur (SPN 190) existe sur DEUX bus (principal et CAN0). En regardant la "
  "DISPONIBILITE du bus CAN0 autour du debut d'attaque (51 episodes) : a ~+4 s, CAN0 chute "
  "brutalement a 0 % (marche d'escalier) pendant que le bus principal reste a 100 %. La "
  "couverture CAN0 passe de 67 % a 6,7 %, dans TOUS les groupes (G1 compris).")
callout("Pourquoi c'est l'injection et non la reaction", "Le silence de CAN0 est un effet DIRECT "
        "de l'injection (l'ELD compromis sature le bus), present meme quand le conducteur ne "
        "reagit pas (Groupe 1). C'est une signature de l'intrusion elle-meme - exactement ce qui "
        "manquait. L'injection etait 'absente en valeur' mais 'presente en disponibilite'.")
figure("v2_injection_signature.png", "Figure 10.1 - Le bus CAN0 se tait ~4 s apres l'onset (signature d'injection).", 6.0)
H("10.2 Caracterisation fine : deux empreintes simultanees", 2)
P("A l'onset, deux signatures coexistent a +4 s : (1) DISPONIBILITE - seul CAN0 se tait "
  "(l'injection) ; (2) VALEUR - une grappe de signaux moteur s'effondre (debit EGR, carburant, "
  "couple, freinage), c'est la REACTION (le vehicule decelere). On separe ainsi, signal par "
  "signal, ce que la cible confond.")
figure("v2_attack_fingerprint.png", "Figure 10.2 - Empreinte de l'attaque : disponibilite (injection) vs valeur (reaction).", 6.0)
P("Limite honnete : ce detecteur de silence CAN0 (PR-AUC 0,53 global) ne sauve pas le Groupe 1, "
  "car son bus CAN0 etait deja peu logue - il n'y a rien a faire taire.")
pagebreak()

# ============================================================================ 11. ROBUSTESSE
H("11. Robustesse et limites (honnetete assumee)", 1)
H("11.1 L'attaquant adaptatif (evasion)", 2)
P("Un attaquant white-box qui connait le modele peut maquiller les signaux surveilles. On "
  "neutralise les top-k signaux (on remet la mediane normale) et on re-score : neutraliser UN "
  "seul signal fait chuter la PR-AUC de 0,74 a 0,22 ; deux signaux a 0,07. Le detecteur est "
  "FRAGILE - son importance est ultra-concentree sur le signal CAN0.")
figure("v2_evasion.png", "Figure 11.1 - Fragilite a un attaquant adaptatif.", 5.5)
H("11.2 Detecteur mono-attaque (taxonomie)", 2)
P("Le dataset ne contient qu'une attaque. On en synthetise d'autres sur des fenetres normales : "
  "le champion attrape l'attaque reelle (65 %) mais RATE le fuzzing, le masquerade furtif et le "
  "replay (~1 %). Perimetre valide = l'attaque ELD du dataset, pas 'les intrusions CAN' en general.")
figure("v2_taxonomy.png", "Figure 11.2 - Generalisation a d'autres attaques (mono-attaque).", 5.5)
H("11.3 Biometrie et clustering / hybride", 2)
P("La fusion biometrie conditionnee par awareness n'aide pas (et degrade le Groupe 1). Le "
  "clustering n'isole pas l'attaque (ARI ~ 0,001) ; le semi-supervise n'ajoute rien ; l'IDS "
  "hybride (champion + signature d'injection) n'apporte pas de gain car le champion absorbe deja "
  "la signature. Autant de negatifs assumes qui renforcent la credibilite.")
figure("v2_hybrid.png", "Figure 11.3 - Clustering, semi-supervise et IDS hybride.", 5.5)
pagebreak()

# ============================================================================ 12. FAISABILITE
H("12. Caracterisation, faisabilite et reproductibilite", 1)
H("12.1 Faisabilite embarquee (mesuree)", 2)
P("Le champion serialise pese 695 Ko ; l'inference est de 2,25 ms par fenetre (562 000 "
  "fenetres/s en lot). Le detecteur est trivial a deployer (passerelle ou ELD). Le goulot n'est "
  "pas le calcul mais la DONNEE : l'agregation a la seconde perd la signature sous-seconde d'une "
  "injection rapide.")
H("12.2 Reproductibilite", 2)
bullets(["Split par conducteur partout ; metrique PR-AUC ; versions verrouillees (requirements.lock).",
         "14 tests unitaires codifient les invariants anti-fuite (aucune feature GPS dans le set CAN).",
         "Deep multi-graine (intervalles de confiance) ; reproductibilite cross-plateforme verifiee "
         "(champion identique Mac/PC GPU)."])
H("12.3 Positionnement vis-a-vis de l'etat de l'art", 2)
P("Le dataset a une publication associee (Lanigan et al., Journal of Transportation Security, "
  "2025) : elle etudie la REPONSE COMPORTEMENTALE du conducteur (temps de reaction 30,3 / 16,1 / "
  "7,5 s selon l'awareness) mais ne construit pas de detecteur. Notre apport est complementaire : "
  "on montre, cote signal CAN, que la detectabilite suit le MEME gradient d'awareness "
  "(0,74 / 0,92 / 0,96). Aucune PR-AUC de detection n'est publiee sur ce dataset - notre cadrage "
  "IDS y semble inedit, d'ou une comparaison au hasard (0,015) et entre nos chemins, faute de SOTA direct.")
table(["Travail", "Apport", "Lien avec notre projet"],
      [["Cho & Shin, CIDS (USENIX Sec. 2016)", "empreinte d'horloge par ECU", "detecterait l'ELD usurpateur independamment du payload ET du conducteur -> reponse a A1"],
       ["Hanselmann, CANet (IEEE Access 2020)", "auto-encodeur par signal", "voie anomalie haute-dimension (la notre echoue en agrege 1 s)"],
       ["Taylor et al. (DSAA 2016)", "LSTM sur sequences CAN", "notre chemin C (GRU) ; meme constat : pas decisif"],
       ["Verma et al., ROAD (ORNL 2022)", "attaques masquerade furtives", "banc d'essai pour la generalisation et l'attaquant adaptatif"],
       ["Murvay & Groza (IEEE T-VT 2018)", "failles de SAE J1939", "fondement de l'absence de securite du bus"]])
callout("Lecture honnete", "Les IDS CAN de reference (CIDS, CANet, ROAD) operent sur la TRAME "
        "brute, pas sur des agregats a la seconde. Notre cadrage, contraint par le dataset deja "
        "agrege, n'est donc pas directement comparable - c'est une limite de portee, pas de methode.")
pagebreak()

# ============================================================================ 13. AUTO-CRITIQUE
H("13. Auto-critique et limites", 1)
P("L'auto-critique fait partie integrante du metier d'AI Engineer.")
H("13.1 Traite au cours du projet", 2)
bullets(["Confondeurs lieu/temps neutralises ; split par conducteur ; PR-AUC.",
         "ROC et optimisation d'hyperparametres ajoutees (demandes du sujet).",
         "Signature d'injection isolee (silence CAN0), distincte de la reaction.",
         "Robustesse : attaquant adaptatif ; taxonomie mono-attaque ; biometrie testee.",
         "Reproductibilite : tests, versions, multi-graine."])
H("13.2 Limites residuelles assumees", 2)
bullets(["Cible conflee : on detecte injection ET reaction ; le label ne les separe pas.",
         "Detecteur evadable par un adversaire qui connait le modele (1-2 signaux).",
         "Mono-attaque : non valide sur fuzzing / replay / masquerade.",
         "Groupe 1 (non averti) : angle mort (ni reaction, ni CAN0 logue).",
         "Generalisation hors-dataset (ROAD/Car-Hacking) non faite (hors scope)."])
H("13.3 Perspectives (au-dela du perimetre du sujet)", 2)
P("Les limites ci-dessus tracent des axes d'amelioration clairs, tous hors du perimetre strict "
  "de la consigne (qui porte sur CE dataset CAN) :")
table(["Limite", "Piste concrete", "Effet attendu"],
      [["Generalisation non prouvee", "rejouer le pipeline sur ROAD / Car-Hacking", "seule vraie preuve de transfert"],
       ["Detecteur evadable", "2e couche ORTHOGONALE (coherence physique inter-signaux ; empreinte d'horloge CIDS)", "augmente le cout d'evasion"],
       ["Cible conflee / Groupe 1", "donnees haute-resolution (1/100 s)", "isoler la signature d'injection sous-seconde"],
       ["Mono-attaque", "entrainer sur une taxonomie (DoS, fuzzing, replay)", "sortir du detecteur mono-signature"]])
P("On a egalement produit une auto-evaluation dans la peau de l'encadrant "
  "(docs/04_conclusion/evaluation_jury.md) : note indicative 17/20, bridee non par la methode "
  "mais par la nature du jeu de donnees (un echantillon, une attaque).")
pagebreak()

# ============================================================================ 14. DEPLOIEMENT / ETHIQUE
H("14. Deploiement, ethique et conduite de projet", 1)
H("14.1 Ou vit l'IDS", 2)
P("Le detecteur se place sur une passerelle filtrante entre l'ELD et le bus critique, ou dans un "
  "module de supervision. Son empreinte (695 Ko, 2 ms) le rend deployable ; le seuil recommande "
  "est HAUTE PRECISION (< 1 % d'alertes) pour limiter la fatigue d'alerte en flotte.")
H("14.2 Ethique, dual-use et RGPD", 2)
P("Projet defensif sur une vulnerabilite reelle (compromission ELD, avis CISA). Les attaques "
  "synthetiques ne servent qu'a l'evaluation. Le dataset contient de la biometrie humaine (HR/EDA) "
  "et de la geolocalisation - donnees personnelles : on a EXCLU le GPS (confondeur) et montre que "
  "la biometrie n'apporte rien, donc un deploiement peut s'en passer (empreinte vie privee reduite).")
H("14.3 Conduite du projet", 2)
P("Phases jalonnees, carnet de bord date, vagues d'amelioration apres auto-critique. Chaque "
  "decision majeure (exclusion GPS, split conducteur, correction d'interpretation) est tracee et "
  "justifiee.")
H("14.4 Le demonstrateur interactif (livrable central)", 2)
P("Au-dela du present rapport, un demonstrateur Streamlit (deliverables/app.py) transforme les "
  "resultats en BANC D'ESSAI : l'evaluateur ne lit pas des chiffres, il manipule le vrai modele et "
  "voit la consequence en direct. La demo ne redit pas les slides - elle apporte l'INTERACTION, "
  "qu'un support fige ne permet pas. Cinq postes, chacun pose une question et y repond par l'action :")
table(["Poste (question)", "Ce qu'on y fait", "Ce qu'on y apprend"],
      [["Voir l'IDS detecter", "rejeu anime d'une attaque reelle", "la detection et le silence du bus CAN0 en temps reel"],
       ["L'IDS est-il robuste ?", "neutraliser les signaux surveilles", "la fragilite : evasion en 1-2 signaux"],
       ["Detecte-t-il l'inconnu ?", "injecter DoS / fuzzing / masquerade", "les angles morts (detecteur mono-attaque)"],
       ["Mon score est-il honnete ?", "choisir le split et les features", "la fuite par conducteur et le confondeur de lieu"],
       ["Combien de fausses alertes ?", "regler seuil et taux d'attaque reel", "la base-rate fallacy et le cout operationnel"]])
callout("Pourquoi une demo plutot qu'une slide de plus", "Un IDS se juge a l'usage : a quel seuil, "
        "avec quels faux positifs, et resiste-t-il a un adversaire. Ces questions se REPONDENT en "
        "manipulant le detecteur - c'est l'objet du demonstrateur, et la transition naturelle vers "
        "la soutenance.")
pagebreak()

# ============================================================================ 15. CONCLUSION
H("15. Conclusion", 1)
P("Partant d'un dataset de flotte sous cyberattaque, nous avons construit un detecteur "
  "d'intrusion complet et, surtout, appris a en douter. Le parcours a mene d'un score apparent a "
  "une comprehension d'ingenieur securite : quel signal est robuste a quoi, a quel cout.")
P("Trois enseignements majeurs : (1) la vraie difficulte etait methodologique (les confondeurs "
  "lieu/temps et la fuite par conducteur), pas le choix du modele ; (2) on peut isoler une "
  "signature de l'injection (silence CAN0) distincte de la reaction du conducteur ; (3) "
  "l'honnetete - resultats negatifs assumes, attaquant adaptatif, base rate realiste - distingue "
  "un resultat credible d'un score flatteur. Le demonstrateur interactif donne corps a l'IDS : un "
  "systeme qui detecte, et reste lucide sur ses incertitudes.")
pagebreak()

H("References", 1)
bullets([
    "Lanigan, Biggs, Gallegos, Daily, Reid, Powers (2025). Impact of Cyber Threat Awareness on "
    "Driver Response to an Unexpected Vehicle Cyberattack. Journal of Transportation Security.",
    "CISA. Advisories sur les vulnerabilites des Electronic Logging Devices (ELD).",
    "Cho K.-T., Shin K. (2016). Fingerprinting ECUs for Vehicle Intrusion Detection (CIDS). USENIX Security.",
    "Hanselmann M. et al. (2020). CANet: An Unsupervised IDS for High-Dimensional CAN Data. IEEE Access.",
    "Verma M. et al. (2022). ROAD: Real ORNL Automotive Dynamometer CAN IDS Dataset.",
    "Murvay P.-S., Groza B. (2018). Security Shortcomings of the SAE J1939 Protocol. IEEE T-VT.",
])
H("Annexe - structure du code et glossaire", 1)
P("Depot : src/ (data, features), notebooks/ (01 a 12, reproductibles), deliverables/ "
  "(demo Streamlit, ce rapport, les slides), docs/ (projet, experiences, evaluation, conclusion), "
  "tests/ (14 tests), artifacts/ (modele entraine).")
P("Glossaire : CAN (bus diffuse), J1939 (protocole poids lourds), SPN (parametre decode), ECU "
  "(calculateur), ELD (boitier reglementaire), IDS (detecteur d'intrusion), PR-AUC (aire "
  "precision-rappel), LODO (leave-one-driver-out), awareness (niveau de prevenance du conducteur).")

doc.save(OUT)
print(f"[OK] rapport -> {OUT}")
print(f"     {len(doc.paragraphs)} paragraphes")
